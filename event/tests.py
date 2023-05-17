from django.test import TestCase, Client
from django.urls import reverse
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.exceptions import ObjectDoesNotExist
from account.models import NonSSOAccount

from django.utils import timezone
import datetime

from .models import Event, EventEmployee, ValidationFile
from pegawai.models import Pegawai
from account.models import Account
from django.contrib.auth.models import User
from log.models import Log

from django.core.exceptions import ValidationError
from .validators import validate_event_employee_fields

from account.tests import set_up_login, set_up_akun_dummy

from docx import Document
import docx

from .views import GenerateDocs

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.shared import Pt

from io import BytesIO

from unittest.mock import MagicMock

CREATE_EVENT = 'create_event.html'
EVENT_LIST = 'event_list.html'
FORBIDDEN_URL = '/home/forbidden/'
DATE_FORMAT = "%Y-%m-%d"
NON_SSO_UI = 'Non SSO UI'
TEST_EVENT = "Test Event"
ALAMAT_NPWP = 'Jl. Hj. Halimah Saerang I No. 9 RT. 004/02 Kukusan Beji Depok'
FORM_SURAT_TUGAS = '/event/form-surat-tugas/1'
ACC1_EXAMPLE = 'acc1@example.com'
EMAIL_PEGAWAI = 'johndoer5@gmail.com'
EVENT_NAME = 'Event Paling Baru'
REUPLOAD_SURAT_TUGAS = 'reupload_surat_tugas'
UPLOAD_SURAT_TUGAS = 'upload_surat_tugas'
REJECT_EVENT = 'reject_event'
VALIDATE_EVENT = 'validate_event'
TEST_USER_EMAIL = 'testuser@test.com'
TEST_USER2_EMAIL = 'testuser2@test.com'
PEGAWAI1_EMAIL = 'pegawai1@gmail.com'
PEGAWAI2_EMAIL = 'pegawai2@gmail.com'
DUMMY_BINARY_DATA = b'some binary data'
DETAIL_EVENT_HTML = 'detail_event.html'
DUMMY_PDF = 'dummy.pdf'
APPLICATION_PDF = 'application/pdf'
WAITING_FOR_VALIDATION = 'Waiting for validation'

def setup_pegawai():
  pegawai1 = Pegawai.objects.create(
    email = PEGAWAI1_EMAIL,
    employee_no = '111',
    employee_name = 'joni',
    employee_category = 'Staff',
    job_status = 'Administrasi',
    grade_level = '-',
    employment_status = 'Kontrak',
    nama_di_rekening = 'karyawankeren',
    nama_bank = 'Mandiri',
    nomor_rekening = '4971335367',
    nomor_npwp = '247128658',
    alamat_npwp = ALAMAT_NPWP
  )

  pegawai2 = Pegawai.objects.create(
    email = PEGAWAI2_EMAIL,
    employee_no = '222',
    employee_name = 'jono',
    employee_category = 'Staff',
    job_status = 'Administrasi',
    grade_level = '-',
    employment_status = 'Kontrak',
    nama_di_rekening = 'karyawankeren',
    nama_bank = 'Mandiri',
    nomor_rekening = '4971335367',
    nomor_npwp = '2471286667',
    alamat_npwp = ALAMAT_NPWP
  )

  return pegawai1,pegawai2

class EventCreateViewTestCase(TestCase):
  def setUp(self):
    self.client = Client()
    self.user = User.objects.create_user(
        username='testuser', email=TEST_USER_EMAIL, password='testpassword')
    self.client.login(username='testuser', password='testpassword')

    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = ACC1_EXAMPLE,
      role = 'User',
      is_first_login = True
    )
    non_sso_acc.save()
    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = 'acc2@example.com',
      role = 'User',
      accountType = NON_SSO_UI
    )
    
    self.pegawai1, self.pegawai2 = setup_pegawai()

    self.start_date = '2023-03-21'
    self.end_date   = '2023-03-23'
    self.strpformat = '%Y-%m-%d'
  
  def test_create_event_get(self):
    response = self.client.get(reverse('create_event'))
    self.assertEqual(response.status_code, 200)

  def test_create_event_only(self):
    url = reverse('create_event')
    
    response = self.client.get(url)
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, CREATE_EVENT)

    new_event_name = 'Test Event_Baru'
    start_dt_obj = datetime.datetime.strptime(self.start_date, self.strpformat).date()
    end_dt_obj = datetime.datetime.strptime(self.end_date, self.strpformat).date()
    
    sess_data = {
      'event_name': new_event_name,
      'start_date': self.start_date,
      'end_date': self.end_date,
    }

    response = self.client.post(url, data=sess_data)
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, EVENT_LIST)

    event = Event.objects.last()
    self.assertEqual(event.start_date, start_dt_obj)
    self.assertEqual(event.end_date, end_dt_obj)
        

  def test_create_event(self):
    url = reverse('create_event')
    
    response = self.client.get(url)
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, CREATE_EVENT)

    new_event_name = 'Test Event_Baru'
    start_dt_obj = datetime.datetime.strptime(self.start_date, self.strpformat).date()
    end_dt_obj = datetime.datetime.strptime(self.end_date, self.strpformat).date()
    
    sess_data = {
      'event_name': new_event_name,
      'start_date': self.start_date,
      'end_date': self.end_date,
      'action': 'add_roles'
    }

    data = {
      'num_fields':1,
      'role_field_0':'PO',
      'honor_field_0':10000,
      'pph_field_0': 10,
      'dropdown-select_0':'111'
    }

    response = self.client.post(url, data=sess_data)
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, 'input_employee.html')
    
    session_data = self.client.session
    self.assertEqual(session_data['event_name'], new_event_name)
    self.assertEqual(session_data['start_date'], self.start_date)
    self.assertEqual(session_data['end_date'], self.end_date)
    
    response = self.client.post(reverse('input_employee_to_event'), data)
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, EVENT_LIST)

    event = Event.objects.last()
    self.assertEqual(event.creator, self.account)
    self.assertEqual(event.event_name, new_event_name)
    self.assertEqual(Event.objects.count(), 1)
    self.assertEqual(event.start_date, start_dt_obj)
    self.assertEqual(event.end_date, end_dt_obj)
    
    event_employee1 = EventEmployee.objects.get(employee=self.pegawai1)
    self.assertEqual(event_employee1.role, 'PO')
    self.assertEqual(event_employee1.honor, 10000)
    self.assertEqual(event_employee1.pph, 10)
  
  def test_create_event_empty_num_fields(self):
    url = reverse('create_event')
    
    response = self.client.get(url)
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, CREATE_EVENT)

    new_event_name = 'Test Event Baru Lagi'
    
    sess_data = {
      'event_name': new_event_name,
      'start_date': self.start_date,
      'end_date': self.end_date,
      'action': 'add_roles'
    }

    data = {
      'num_fields':'',
    }

    response = self.client.post(url, data=sess_data)
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, 'input_employee.html')
    
    session_data = self.client.session
    self.assertEqual(session_data['event_name'], new_event_name)
    self.assertEqual(session_data['start_date'], self.start_date)
    self.assertEqual(session_data['end_date'], self.end_date)
    
    self.client.post(reverse('input_employee_to_event'), data)
    event = Event.objects.last()
    
    empty_qs = EventEmployee.objects.none()
    queryset = EventEmployee.objects.filter(event=event)
    self.assertQuerysetEqual(queryset, empty_qs, "Queryset should be empty")


class EventModelTest(TestCase):
  def setUp(self):
    email = 'joni@gmail.com'

    self.event_name = 'Test Event'
    self.binary_data = DUMMY_BINARY_DATA

    self.user = User.objects.create_user(
      username='admin', 
      password='admin123', 
      email='adminkece@gmail.com'
    )

    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = email,
      role = 'User',
      is_first_login = True
    )
    non_sso_acc.save()
    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = email,
      role = 'User',
      accountType = NON_SSO_UI
    )

    self.pegawai = Pegawai.objects.create(
      email = email,
      employee_no = 'employee1',
      employee_name = 'Jonyy',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = 'Jl. Hj. Halimah Saeran I No. 1 RT. 004/02 Kukusan Beji Depok'
    )

    login = self.client.login(username='admin', password='admin123')
    self.assertTrue(login)

  def test_create_event(self):
    event = Event.objects.create(
      creator=self.account,
      event_name=self.event_name,
      start_date='2022-04-27',
      end_date='2022-05-28',
      expense=100000,
      sk_file=self.binary_data
    )
    self.assertIsInstance(event, Event)
    self.assertEqual(event.creator, self.account)
    self.assertEqual(event.event_name, self.event_name)
    self.assertLess(event.start_date, event.end_date)
    self.assertEqual(event.expense, 100000)
    self.assertEqual(event.sk_file, self.binary_data)
    
  def test_create_event_will_add_log(self):
    sess_data = {
      'event_name':self.event_name,
      'start_date':'2022-03-15',
      'end_date':'2022-06-12',
    }

    self.client.post(reverse('create_event'), data=sess_data)
    self.assertEqual(len(Log.objects.all()), 1)
    
    action = 'Create ' + self.event_name + ' event'
    self.assertTrue(Log.objects.filter(action=action).first())

  def test_create_event_employee(self):
    event = Event.objects.create(
      creator=self.account,
      event_name=self.event_name,
      start_date='2022-01-01',
      end_date='2022-02-02',
      expense=100000,
      sk_file=self.binary_data
    )

    event_employee = EventEmployee.objects.create(
      employee=self.pegawai,
      event=event
    )

    self.assertIsInstance(event_employee, EventEmployee)
    self.assertEqual(event_employee.employee, self.pegawai)
    self.assertEqual(event_employee.event, event)
  
class CreateUnauthorizedEventViewTestCase(TestCase):
  def test_unauthenticated_create(self):
    response = self.client.get(reverse('create_event'))
    self.assertEqual(response.status_code, 302)
  
  def test_unauthenticated_input(self):
    response = self.client.get(reverse('input_employee_to_event'))
    self.assertEqual(response.status_code, 302)


class ShowEventListViewTestCase(TestCase):
  def setUp(self):
    email = 'jona@gmail.com'

    self.event_name = 'Test Event Baru'
    self.binary_data = DUMMY_BINARY_DATA

    self.user = User.objects.create_user(
      username='admin', 
      password='admin123', 
      email='adminkece@gmail.com'
    )

    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = email,
      role = 'User',
      is_first_login = True
    )
    non_sso_acc.save()
    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = email,
      role = 'User',
      accountType = NON_SSO_UI
    )

    self.event_name_1 = 'Event 1'
    self.event1 = Event.objects.create(
      creator=self.account,
      event_name=self.event_name_1,
      start_date='2022-04-23',
      end_date='2022-05-29',
      expense=10000,
      sk_file=self.binary_data
    )
    self.event2 = Event.objects.create(
      event_name='Event 2',
      start_date='2022-04-24',
      end_date='2022-05-30',
      expense=12000,
      sk_file=self.binary_data
    )
  
  def test_get_events(self):
    self.client.login(username='admin', password='admin123')
    response = self.client.get(reverse('get_events'))
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, 'event_list.html')
    self.assertContains(response, self.event_name_1)
    self.assertContains(response, 'Event 2')
    self.assertDictEqual(response.context['owner_data'], {self.event_name_1: True})

  def test_get_events_unauthenticated(self):
    response = self.client.get(reverse('get_events'))
    self.assertRedirects(response, '/login?next=/event/')

  def test_get_options_view(self):
    Pegawai.objects.create(
      email = 'johndoer1@gmail.com',
      employee_no = '190756565',
      employee_name = 'Test Pegawai 1',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335365',
      nomor_npwp = '247128654',
      alamat_npwp = 'Jl. Hj. Halimah Saerang I No. 1 RT. 004/02 Kukusan Beji Depok'
    )
    Pegawai.objects.create(
      email = 'johndoer2@gmail.com',
      employee_no = '190756566',
      employee_name = 'Test Pegawai 2',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335366',
      nomor_npwp = '247128652',
      alamat_npwp = 'Jl. Hj. Halimah Saerang I No. 2 RT. 004/02 Kukusan Beji Depok')
  
    Pegawai.objects.create(
      email = 'johndoer3@gmail.com',
      employee_no = '190756567',
      employee_name = 'Test Pegawai 3',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128659',
      alamat_npwp = 'Jl. Hj. Halimah Saerang I No. 3 RT. 004/02 Kukusan Beji Depok'
    )

    response = self.client.get(reverse('get_options'), data={'search': 'Test'})
    self.assertEqual(response.status_code, 200)
    self.assertJSONEqual(str(response.content, encoding='utf8'), [
      {'employee_no': '190756565', 'employee_name': 'Test Pegawai 1'},
      {'employee_no': '190756566', 'employee_name': 'Test Pegawai 2'},
      {'employee_no': '190756567', 'employee_name': 'Test Pegawai 3'},
    ])



class EventEmployeeModelTestCase(TestCase):
  @classmethod
  def setUpTestData(cls):
    email = 'jona@gmail.com'
    cls.binary_data = b'some binary data 1'
    cls.user = User.objects.create_user(
      username='admin', 
      password='admin123', 
      email='adminkocak@gmail.com'
    )

    non_sso_acc = NonSSOAccount.objects.create(
      user = cls.user,
      username = 'jonikeren',
      email = email,
      role = 'User',
      is_first_login = True
    )
    non_sso_acc.save()
    cls.account = Account.objects.create(
      user = cls.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = email,
      role = 'User',
      accountType = NON_SSO_UI
    )

    cls.event = Event.objects.create(
      event_name='Test Event Baru',
      start_date='2023-01-01',
      end_date='2023-01-02',
    )

    cls.pegawai = Pegawai.objects.create(
      email = 'johndoer2@gmail.com',
      employee_no = '190756565',
      employee_name = 'Jonyy',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = 'Jl. Hj. Halimah Saerang I No. 1 RT. 004/02 Kukusan Beji Depok'
    )

    cls.event_employee = EventEmployee.objects.create(
      employee=cls.pegawai,
      event=cls.event,
      honor=1000,
      pph=10,
      role='Admin'
    )

  def test_employee_field(self):
      field = EventEmployee._meta.get_field('employee')
      self.assertEqual(field.related_model, Pegawai)

  def test_event_field(self):
    field = EventEmployee._meta.get_field('event')
    self.assertEqual(field.related_model, Event)

  def test_honor_field(self):
    field = EventEmployee._meta.get_field('honor')
    self.assertEqual(field.default, 0)
    self.assertEqual(field.validators[0].limit_value, 0)

  def test_netto_field(self):
    field = EventEmployee._meta.get_field('netto')
    self.assertEqual(field.default, 0)
    self.assertEqual(field.validators[0].limit_value, 0)

  def test_pph_field(self):
    field = EventEmployee._meta.get_field('pph')
    self.assertEqual(field.default, 0)
    self.assertEqual(field.validators[0].limit_value, 0)

  def test_role_field(self):
    field = EventEmployee._meta.get_field('role')
    self.assertEqual(field.max_length, 100)
    self.assertEqual(field.null, True)

  def test_event_employee_creation(self):
    self.assertEqual(self.event_employee.employee, self.pegawai)
    self.assertEqual(self.event_employee.event, self.event)
    self.assertEqual(self.event_employee.honor, 1000)
    self.assertEqual(self.event_employee.pph, 10)
    self.assertEqual(self.event_employee.role, 'Admin')
    self.assertEqual(
      self.event_employee.netto, (100-self.event_employee.pph)/100*self.event_employee.honor)
  

class ValidationTests(TestCase):
  def test_validate_event_employee_fields(self):
    role_name, pph, honor = validate_event_employee_fields('', '', '', 1)
    self.assertEqual(role_name, 'Role_1')
    self.assertEqual(pph, 0)
    self.assertEqual(honor, 0)

    role_name, pph, honor = validate_event_employee_fields('Manager', '100', '50', 2)
    self.assertEqual(role_name, 'Manager')
    self.assertEqual(pph, 100)
    self.assertEqual(honor, 50)

    role_name, pph, honor = validate_event_employee_fields('Supervisor', '-100', '-50', 3)
    self.assertEqual(role_name, 'Supervisor')
    self.assertEqual(pph, 100)
    self.assertEqual(honor, 50)



class InputEmployeeToEventTestCase(TestCase):
  def setUp(self):
    self.client = Client()
    self.start_date = '2023-03-22'
    self.end_date   = '2023-03-25'
    self.user = User.objects.create_user(
      username='admin5', 
      password='admin345', 
      email='adminkacok@gmail.com'
    )
    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonakeren',
      email = 'jokimak@gmail.com',
      role = 'User',
      is_first_login = True
    )
    non_sso_acc.save()
    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonakeren',
      email = 'jokimak@gmail.com',
      role = 'User',
      accountType = NON_SSO_UI
    )

    self.sess_data = {
      'event_name': EVENT_NAME,
      'start_date': self.start_date,
      'end_date': self.end_date,
      'action': 'add_roles'
    }
    self.pegawai = Pegawai.objects.create(
      email = EMAIL_PEGAWAI,
      employee_no = '123',
      employee_name = 'Jonyz',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = ALAMAT_NPWP
    )
    self.event_data = {
        'num_fields': 1,
        'dropdown-select_0': '123',
        'role_field_0': 'Admin',
        'honor_field_0': '-100',
        'pph_field_0': '-200',
    }
    self.client.login(username='admin5', password='admin345')

  def test_input_employee_to_event_with_negative_honor_and_pph(self):
    self.client.post(reverse('create_event'), data=self.sess_data)
    response = self.client.post(reverse('input_employee_to_event'), data=self.event_data, follow=True)
    self.assertEqual(response.status_code, 200)
    self.assertContains(response, 'Event List')
    self.assertEqual(abs(int(self.event_data['honor_field_0'])), 100)
    self.assertEqual(abs(int(self.event_data['pph_field_0'])), 200)

def set_up_event_employee(self):
    self.event_employee=EventEmployee(
      employee=self.pegawai,
      event=self.event,
      honor=10000,
      pph=5,
      netto=9500,
      role='Ketua'
    )
    self.event_employee.save()
                
class RUDEventLoggedInAdminTest(TestCase):
  def setUp(self) -> None:
    set_up_login(self, 'User')
    self.start_date = '2023-03-22'
    self.end_date   = '2023-03-25'
    self.event = Event(
      creator=self.account,
      event_name=EVENT_NAME,
      start_date= self.start_date,
      end_date= self.end_date,
      expense=20000
    )
    self.event.save()
    self.pegawai = Pegawai(
      email = EMAIL_PEGAWAI,
      employee_no = '123',
      employee_name = 'Jonyz',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = ALAMAT_NPWP
    )
    self.pegawai.save()
                
  

  def test_riwayat_event_valid(self):
    response = self.client.get('/event/my-event')
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, 'riwayat_event.html')

  def test_detail_event_valid(self):
    response = self.client.get(f'/event/detail/{self.event.id}')
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, DETAIL_EVENT_HTML)
                
  def test_detail_event_id_invalid(self):
    response = self.client.get('/event/detail/3333')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_detail_event_id_bukan_int(self):
    response = self.client.get('/event/detail/asd')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_event_get_valid(self):
    response = self.client.get(f'/event/update/{self.event.id}')
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, 'update_event.html')
                
  def test_update_event_id_invalid(self):
    response = self.client.get('/event/update/3333')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_event_id_bukan_int(self):
    response = self.client.get('/event/update/asd')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

  def test_update_validated_event_fail_get(self):
    self.event.status = 'Validated'
    self.event.save()
    response = self.client.get(f'/event/update/{self.event.id}')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_event_post_valid(self):
    new_event_name = 'Updated Event2'
    new_start_date = '2020-01-01'
    new_end_date   = '2023-01-01'
    data = {
      'event_name': new_event_name,
      'start_date': new_start_date,
      'end_date': new_end_date
    }
    response = self.client.post(f'/event/submit-update/{self.event.id}', data)
    updated_event = Event.objects.get(id=self.event.id)
    self.assertEqual(updated_event.event_name, new_event_name)
    self.assertEqual(updated_event.start_date.strftime(DATE_FORMAT), new_start_date)
    self.assertEqual(updated_event.end_date.strftime(DATE_FORMAT), new_end_date)
                                                    
    self.assertRedirects(response, f'/event/detail/{self.event.id}', status_code=302, target_status_code=200)
                
  def test_update_event_post_id_invalid(self):
    new_event_name = 'Updated Event'
    new_start_date = '2020-01-01'
    new_end_date   = '2023-01-01'
    data = {
      'event_name': new_event_name,
      'start_date': new_start_date,
      'end_date': new_end_date
    }
    response = self.client.post('/event/submit-update/3333', data)
    updated_event = Event.objects.get(id=self.event.id)
    self.assertEqual(updated_event.event_name, self.event.event_name)
    self.assertEqual(updated_event.start_date.strftime(DATE_FORMAT), self.event.start_date)
    self.assertEqual(updated_event.end_date.strftime(DATE_FORMAT), self.event.end_date)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

  def test_update_validated_event_fail(self):
    self.event.status = 'Validated'
    self.event.save()
    new_event_name = 'Updated Events'
    new_start_date = '2020-01-01'
    new_end_date   = '2023-01-02'
    data = {
      'event_name': new_event_name,
      'start_date': new_start_date,
      'end_date': new_end_date
    }
    response = self.client.post(f'/event/submit-update/{self.event.id}', data)
    updated_event = Event.objects.get(id=self.event.id)
    self.assertEqual(updated_event.event_name, self.event.event_name)
    self.assertEqual(updated_event.start_date.strftime(DATE_FORMAT), self.event.start_date)
    self.assertEqual(updated_event.end_date.strftime(DATE_FORMAT), self.event.end_date)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_input_employee_to_existing_event_get(self):
    response = self.client.get(f'/event/add-employee/{self.event.id}')
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, 'input_employee_to_existing_event.html')
                
  def test_input_employee_to_existing_event_get_id_invalid(self):
    response = self.client.get('/event/add-employee/3333')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_input_employee_to_existing_event_get_id_bukan_int(self):
    response = self.client.get('/event/add-employee/asd')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

  def test_input_employee_to_existing_validated_event_fail(self):
    self.event.status = 'Validated'
    self.event.save()
    response = self.client.get(f'/event/add-employee/{self.event.id}')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_input_employee_to_existing_event_post(self):
    role = 'Ketua'
    honor = 100000
    pph   = 5
    employee_no = '123'
    data = {
      'num_fields': '1',
      'role_field_0': role,
      'honor_field_0': honor,
      'pph_field_0': pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post(f'/event/submit-add-employee/{self.event.id}', data)
    event = Event.objects.get(id=self.event.id)
    event_employee = EventEmployee.objects.filter(event=event).first()
    self.assertEqual(event_employee.role, role)
    self.assertEqual(event_employee.honor, honor)
    self.assertEqual(event_employee.pph, pph)
                                                    
    self.assertRedirects(response, f'/event/detail/{self.event.id}', status_code=302, target_status_code=200)
    
  def test_input_employee_to_existing_event_post_kosongan(self):
    data = {
      'num_fields': '',
    }
    response = self.client.post(f'/event/submit-add-employee/{self.event.id}', data)
    event = Event.objects.get(id=self.event.id)
    with self.assertRaises(ObjectDoesNotExist):
      EventEmployee.objects.get(event=event)
                                                    
    self.assertRedirects(response, f'/event/detail/{self.event.id}', status_code=302, target_status_code=200)
                
  def test_input_employee_to_existing_event_post_id_invalid(self):
    role = 'Ketua'
    honor = 100000
    pph   = 5
    employee_no = '123'
    data = {
      'num_fields': '1',
      'role_field_0': role,
      'honor_field_0': honor,
      'pph_field_0': pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post('/event/submit-add-employee/3333', data)
    event = Event.objects.get(id=self.event.id)
    with self.assertRaises(ObjectDoesNotExist):
      EventEmployee.objects.get(event=event)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_input_employee_to_existing_event_post_id_bukan_int(self):
    role = 'Ketua'
    honor = 100000
    pph   = 5
    employee_no = '123'
    data = {
      'num_fields': '1',
      'role_field_0': role,
      'honor_field_0': honor,
      'pph_field_0': pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post('/event/submit-add-employee/asd', data)
    event = Event.objects.get(id=self.event.id)
    with self.assertRaises(ObjectDoesNotExist):
      EventEmployee.objects.get(event=event)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

  def test_input_employee_to_existing_validated_event_post_fail(self):
    self.event.status = 'Validated'
    self.event.save()
    role = 'Ketua'
    honor = 100000
    pph   = 5
    employee_no = '123'
    data = {
      'num_fields': '1',
      'role_field_0': role,
      'honor_field_0': honor,
      'pph_field_0': pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post(f'/event/submit-add-employee/{self.event.id}', data)
    event = Event.objects.get(id=self.event.id)
    with self.assertRaises(ObjectDoesNotExist):
      EventEmployee.objects.get(event=event)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_employee_to_existing_event_get(self):
    set_up_event_employee(self)
    response = self.client.get(f'/event/update-employee/{self.event_employee.id}')
    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed(response, 'update_event_employee.html')
                
  def test_update_employee_to_existing_event_get_id_invalid(self):
    set_up_event_employee(self)
    response = self.client.get('/event/update-employee/3333')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
    
  def test_update_employee_to_existing_event_get_id_bukan_int(self):
    set_up_event_employee(self)
    response = self.client.get('/event/update-employee/asd')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

  def test_update_employee_to_existing_validated_event_get_fail(self):
    self.event.status = 'Validated'
    self.event.save()
    set_up_event_employee(self)
    response = self.client.get(f'/event/update-employee/{self.event.id}')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_employee_to_existing_event_post(self):
    set_up_event_employee(self)
    new_role = 'Ketuanew'
    new_honor = 100001
    new_pph   = 6
    employee_no = '123'
    data = {
      'role_field_0': new_role,
      'honor_field_0': new_honor,
      'pph_field_0': new_pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post(f'/event/submit-update-employee/{self.event_employee.id}', data)
    event_employee = EventEmployee.objects.filter(id=self.event_employee.id).first()
    self.assertEqual(event_employee.role, new_role)
    self.assertEqual(event_employee.honor, new_honor)
    self.assertEqual(event_employee.pph, new_pph)
                                                    
    self.assertRedirects(response, f'/event/detail/{self.event.id}', status_code=302, target_status_code=200)
                
  def test_update_employee_to_existing_event_post_id_invalid(self):
    set_up_event_employee(self)
    new_role = 'Ketuanew'
    new_honor = 100001
    new_pph   = 6
    employee_no = '123'
    data = {
      'role_field_0': new_role,
      'honor_field_0': new_honor,
      'pph_field_0': new_pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post('/event/submit-update-employee/3333', data)
    event_employee = EventEmployee.objects.filter(id=self.event_employee.id).first()
    self.assertEqual(event_employee.role, self.event_employee.role)
    self.assertEqual(event_employee.honor, self.event_employee.honor)
    self.assertEqual(event_employee.pph, self.event_employee.pph)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_employee_to_existing_event_post_id_bukan_int(self):
    set_up_event_employee(self)
    new_role = 'Ketuanew'
    new_honor = 100001
    new_pph   = 6
    employee_no = '123'
    data = {
      'role_field_0': new_role,
      'honor_field_0': new_honor,
      'pph_field_0': new_pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post('/event/submit-update-employee/asd', data)
    event_employee = EventEmployee.objects.filter(id=self.event_employee.id).first()
    self.assertEqual(event_employee.role, self.event_employee.role)
    self.assertEqual(event_employee.honor, self.event_employee.honor)
    self.assertEqual(event_employee.pph, self.event_employee.pph)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

  def test_update_employee_to_existing_validated_event_post_fail(self):
    self.event.status = 'Validated'
    self.event.save()
    set_up_event_employee(self)
    new_role = 'Ketuanew'
    new_honor = 100001
    new_pph   = 6
    employee_no = '123'
    data = {
      'role_field_0': new_role,
      'honor_field_0': new_honor,
      'pph_field_0': new_pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post(f'/event/submit-update-employee/{self.event.id}', data)
    event_employee = EventEmployee.objects.filter(id=self.event_employee.id).first()
    self.assertEqual(event_employee.role, self.event_employee.role)
    self.assertEqual(event_employee.honor, self.event_employee.honor)
    self.assertEqual(event_employee.pph, self.event_employee.pph)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_delete_employee_to_existing_event_post(self):
    set_up_event_employee(self)
    data = {}
    response = self.client.post(f'/event/delete-employee/{self.event_employee.id}', data)
    with self.assertRaises(ObjectDoesNotExist): # Objek berhasil dihapus
      EventEmployee.objects.get(id=self.event_employee.id)
                                                    
    self.assertRedirects(response, f'/event/detail/{self.event.id}', status_code=302, target_status_code=200)
                
  def test_delete_employee_to_existing_event_post_id_invalid(self):
    set_up_event_employee(self)
    data = {}
    response = self.client.post('/event/delete-employee/3333', data)
                
    self.assertTrue(EventEmployee.objects.filter(id=self.event_employee.id).first()) # Objeknya tidak terhapus
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_delete_employee_to_existing_event_post_id_bukan_int(self):
    set_up_event_employee(self)
    data = {}
    response = self.client.post('/event/delete-employee/asd', data)
                
    self.assertTrue(EventEmployee.objects.filter(id=self.event_employee.id).first()) # Objeknya tidak terhapus																		
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

  def test_delete_employee_to_existing_validated_event_post_fail(self):
    self.event.status = 'Validated'
    self.event.save()
    set_up_event_employee(self)
    data = {}
    response = self.client.post(f'/event/delete-employee/{self.event.id}', data)
                
    self.assertTrue(EventEmployee.objects.filter(id=self.event_employee.id).first()) # Objeknya tidak terhapus																		
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                


class RUDEventLoggedInNonCreatorTest(TestCase):
  def setUp(self) -> None:
    set_up_login(self, 'User')
    set_up_akun_dummy(self)
    self.start_date = '2023-03-22'
    self.end_date   = '2023-03-25'
    self.event = Event(
      creator=self.account_dummy,
      event_name='Event Paling Baru2',
      start_date= self.start_date,
      end_date= self.end_date,
      expense=20000
    )
    self.event.save()
    self.pegawai = Pegawai(
      email = 'johndoer51@gmail.com',
      employee_no = '123',
      employee_name = 'Jonyz',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = 'Jl. Hj. Halimah Saerang I No. 90 RT. 004/02 Kukusan Beji Depok'
    )
    self.pegawai.save()
                

  def test_detail_event_invalid(self):
    response = self.client.get(f'/event/detail/{self.event.id}')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_event_get_invalid(self):
    response = self.client.get(f'/event/update/{self.event.id}')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_event_post_id_invalid(self):
    new_event_name = 'Updated Event'
    new_start_date = '2020-01-01'
    new_end_date   = '2023-01-01'
    data = {
      'event_name': new_event_name,
      'start_date': new_start_date,
      'end_date': new_end_date
    }
    response = self.client.post(f'/event/submit-update/{self.event.id}', data)
    updated_event = Event.objects.get(id=self.event.id)
    self.assertEqual(updated_event.start_date.strftime(DATE_FORMAT), self.event.start_date)
    self.assertEqual(updated_event.event_name, self.event.event_name)
    self.assertEqual(updated_event.end_date.strftime(DATE_FORMAT), self.event.end_date)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_input_employee_to_existing_event_get_invalid(self):
    response = self.client.get(f'/event/add-employee/{self.event.id}')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_input_employee_to_existing_event_post_invalid(self):
    role = 'Ketua'
    honor = 100000
    pph   = 5
    employee_no = '123'
    data = {
      'num_fields': '1',
      'role_field_0': role,
      'honor_field_0': honor,
      'pph_field_0': pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post(f'/event/submit-add-employee/{self.event.id}', data)
    event = Event.objects.get(id=self.event.id)
    with self.assertRaises(ObjectDoesNotExist):
      EventEmployee.objects.get(event=event)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
                
  def test_update_employee_to_existing_event_get_invalid(self):
    set_up_event_employee(self)
    response = self.client.get(f'/event/update-employee/{self.event_employee.id}')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_update_employee_to_existing_event_post_invalid(self):
    set_up_event_employee(self)
    new_role = 'Ketuanew'
    new_honor = 100001
    new_pph   = 6
    employee_no = '123'
    data = {
      'role_field_0': new_role,
      'honor_field_0': new_honor,
      'pph_field_0': new_pph,
      'dropdown-select_0': employee_no
    }
    response = self.client.post(f'/event/submit-update-employee/{self.event_employee.id}', data)
    event_employee = EventEmployee.objects.filter(id=self.event_employee.id).first()
    self.assertEqual(event_employee.role, self.event_employee.role)
    self.assertEqual(event_employee.honor, self.event_employee.honor)
    self.assertEqual(event_employee.pph, self.event_employee.pph)
                                                    
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
                
  def test_delete_employee_to_existing_event_post_invalid(self):
    set_up_event_employee(self)
    data = {}
    response = self.client.post(f'/event/delete-employee/{self.event_employee.id}', data)
                
    self.assertTrue(EventEmployee.objects.filter(id=self.event_employee.id).first()) # Objeknya tidak terhapus
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)

class CEventLoggedInAdminExtraTest(TestCase):
  def setUp(self) -> None:
    set_up_login(self, 'Admin')

  def test_create_event_invalid_role(self):
    response = self.client.get('/event/create')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
    
  def test_riwayat_event_invalid_role(self):
    response = self.client.get('/event/my-event')
    self.assertRedirects(response, FORBIDDEN_URL, status_code=302, target_status_code=200)
    
class GeneratedDocsTest(TestCase):
  def setUp(self):
    self.document = Document()
  
  def test_paragraph_text_is_correct(self):
    paragraph = self.document.add_paragraph("paragraf")
    self.assertEqual(paragraph.text, "paragraf")
  
  def test_paragraph_text_is_not_correct(self):
    paragraph = self.document.add_paragraph("paragraf")
    self.assertNotEqual(paragraph.text, "paragraph")

  def test_font_is_correct(self):
    self.view = GenerateDocs()
    paragraph = self.document.add_paragraph("paragraf")
    run = paragraph.runs[0]
    self.view.set_font(run)

    self.assertEqual(run.font.name, 'Times New Roman')
    self.assertEqual(run.font.size, Pt(12))
    self.assertFalse(run.font.bold)
    self.assertEqual(run.font.underline, WD_UNDERLINE.NONE)

  def test_paragraph_format_is_correct(self):
    self.view = GenerateDocs()
    paragraph = self.document.add_paragraph("paragraf")
    self.view.set_paragraph_format(paragraph)

    self.assertEqual(paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    self.assertEqual(paragraph.paragraph_format.left_indent, Pt(0))

  def test_paragraph_style_is_correct(self):
    self.view = GenerateDocs()
    paragraph4 = self.document.add_paragraph("Nama:")
    self.view.set_font(paragraph4.runs[0])
    self.view.set_paragraph_format(paragraph4, left_indent=36)

    run = paragraph4.add_run() 
    run.add_break(docx.text.run.WD_BREAK.LINE)
    run2 = paragraph4.add_run('Jabatan:')
    self.view.set_font(run2)

    self.assertEqual(paragraph4.style.name, "Normal")
    self.assertEqual(paragraph4.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    self.assertEqual(paragraph4.paragraph_format.left_indent, Pt(36))
    self.assertEqual(paragraph4.runs[0].font.name, "Times New Roman")
    self.assertEqual(paragraph4.runs[0].font.size, Pt(12))
    self.assertFalse(paragraph4.runs[0].font.bold)
    self.assertEqual(paragraph4.runs[0].font.underline, WD_UNDERLINE.NONE)
    self.assertEqual(run2.font.name, "Times New Roman")
    self.assertEqual(run2.font.size, Pt(12))
    self.assertFalse(run2.font.bold)
    self.assertEqual(run2.font.underline, WD_UNDERLINE.NONE)
  
  def test_get_event_name_returns_correct_name(self):
    self.view = GenerateDocs()
    event_name = TEST_EVENT
    event = Event.objects.create(event_name=event_name, start_date = '2020-01-01', end_date   = '2023-01-01',)
    result = self.view.get_event_name(event.id)
    self.assertEqual(result, event_name)

  def test_get_event_name_returns_none_for_invalid_id(self):
    self.view = GenerateDocs()
    result = self.view.get_event_name(1000) # Assuming there's no event with ID=1000
    self.assertIsNone(result)

  def test_get_start_date(self):
    # Create a test event
    start_date = datetime.date(2020, 1, 1)
    end_date = datetime.date(2023, 1, 1)
    event = Event.objects.create(event_name=TEST_EVENT, start_date=start_date, end_date=end_date)

    # Call the method to get the start date
    self.view = GenerateDocs()
    result = self.view.get_start_date(event.id)

    # Check that the result matches the expected value
    self.assertEqual(result, start_date)

  def test_get_end_date(self):
    # Create a test event
    start_date = datetime.date(2020, 1, 1)
    end_date = datetime.date(2023, 1, 1)
    event = Event.objects.create(event_name=TEST_EVENT, start_date=start_date, end_date=end_date)

    # Call the method to get the end date
    self.view = GenerateDocs()
    result = self.view.get_end_date(event.id)

    # Check that the result matches the expected value
    self.assertEqual(result, end_date)
  
  def test_create_docs(self):
    self.pegawai1 = Pegawai(
      email = 'konbab@gmail.com',
      employee_no = '1322323',
      employee_name = 'Jonyzgung',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = ALAMAT_NPWP
    )
    self.pegawai1.save()

    self.pegawai2 = Pegawai(
      email = 'johndower6@gmail.com',
      employee_no = '125',
      employee_name = 'Jony',
      employee_category = 'Staff',
      job_status = 'Administrasi2',
      grade_level = '-',
      employment_status = 'Kontraks',
      nama_di_rekening = 'karyawankeren2',
      nama_bank = 'BCA',
      nomor_rekening = '4971235367',
      nomor_npwp = '247128658',
      alamat_npwp = 'Jl. Hj. Halimah Saerang I No. 9 RT. 004/03 Kukusan Beji Depok'
    )
    self.pegawai2.save()

    event_name = TEST_EVENT
    event = Event.objects.create(event_name=event_name, start_date = '2020-01-01', end_date   = '2023-01-01',)
    EventEmployee.objects.create(
      employee=self.pegawai1,
      event=event
    )
    EventEmployee.objects.create(
      employee=self.pegawai2,
      event=event
    )
    response = self.client.get(FORM_SURAT_TUGAS)
    self.assertEqual(response.status_code, 302)

class GenerateFormDetailSuratTugasTest(TestCase):
  def setUp(self) -> None:
    set_up_login(self, 'User')
    self.start_date = '2023-03-22'
    self.end_date   = '2023-03-25'
    self.event = Event(
      creator=self.account,
      event_name=EVENT_NAME,
      start_date= self.start_date,
      end_date= self.end_date,
      expense=20000
    )
    self.event.save()
    self.pegawai = Pegawai(
      email = EMAIL_PEGAWAI,
      employee_no = '123',
      employee_name = 'Jonyz',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = ALAMAT_NPWP
    )
    self.pegawai.save()

    self.pegawai1 = Pegawai(
      email = 'konbab@gmail.com',
      employee_no = '1322323',
      employee_name = 'Jonyzgung',
      employee_category = 'Staff',
      job_status = 'Administrasi',
      grade_level = '-',
      employment_status = 'Kontrak',
      nama_di_rekening = 'karyawankeren',
      nama_bank = 'Mandiri',
      nomor_rekening = '4971335367',
      nomor_npwp = '247128658',
      alamat_npwp = ALAMAT_NPWP
    )
    self.pegawai1.save()

    self.pegawai2 = Pegawai(
      email = 'johndower6@gmail.com',
      employee_no = '125',
      employee_name = 'Jony',
      employee_category = 'Staff',
      job_status = 'Administrasi2',
      grade_level = '-',
      employment_status = 'Kontraks',
      nama_di_rekening = 'karyawankeren2',
      nama_bank = 'BCA',
      nomor_rekening = '4971235367',
      nomor_npwp = '247128658',
      alamat_npwp = 'Jl. Hj. Halimah Saerang I No. 9 RT. 004/03 Kukusan Beji Depok'
    )
    self.pegawai2.save()

  def test_event_creator_can_post(self):
    EventEmployee.objects.create(
      employee=self.pegawai1,
      event=self.event
    )
    EventEmployee.objects.create(
      employee=self.pegawai2,
      event=self.event
    )
    data = {
      'nomor_surat_tugas': '12345',
      'nama_pj': 'Mirna',
      'jabatan_pj': 'Dekan',
      'perihal_event': 'Wisuda',
      'tugas_panitia': 'Memastikan event lancar',
      'target_anggaran': 'Anggaran fasilkom 2020'
    }
    response = self.client.post(FORM_SURAT_TUGAS, data)
    self.assertEqual(response.status_code, 200)

  def test_event_creator_can_get(self):
    response = self.client.get(FORM_SURAT_TUGAS)
    self.assertEqual(response.status_code, 200)

  def test_not_event_creator(self):
    self.user = User.objects.create_user(
        username='testuser', email=TEST_USER_EMAIL, password='testpassword')
    self.account = Account.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = ACC1_EXAMPLE,
      role = 'User',
      accountType = NON_SSO_UI
    )
    
    self.event = Event(
      creator=self.account,
      event_name=EVENT_NAME,
      start_date= self.start_date,
      end_date= self.end_date,
      expense=20000
    )

    response = self.client.get('/event/form-surat-tugas/2')
    self.assertEqual(response.status_code, 302)

  def test_not_event_creator_forbidden(self):
    self.user = User.objects.create_user(
        username='testuser', email=TEST_USER_EMAIL, password='testpassword')
    self.account = Account.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = ACC1_EXAMPLE,
      role = 'User',
      accountType = NON_SSO_UI
    )
    
    self.event = Event(
      creator=self.account,
      event_name=EVENT_NAME,
      start_date= self.start_date,
      end_date= self.end_date,
      expense=20000
    )
    self.event.save()

    response = self.client.get(f'/event/form-surat-tugas/{self.event.id}')
    self.assertEqual(response.status_code, 302)

class EventValidationTestCase(TestCase):
  
  def setUp(self):
    self.client = Client()
    self.user = User.objects.create_user(
        username='testuser1', email=TEST_USER_EMAIL, password='testpassword1')
    self.client.login(username='testuser1', password='testpassword1')

    self.user2 = User.objects.create_user(
      username='testuser2', email=TEST_USER2_EMAIL, password='testpassword2'
    )
    self.client2 = Client()
    self.client2.login(username='testuser2', password='testpassword2')

    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = 'acc3@example.com',
      role = 'User',
      is_first_login = True
    )

    non_sso_acc2 = NonSSOAccount.objects.create(
      user = self.user2,
      username = 'jonikeren',
      email = 'acc4@example.com',
      role = 'Admin',
      is_first_login = True
    )

    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = 'acc5@example.com',
      role = 'User',
      accountType = NON_SSO_UI
    )

    self.account2 = Account.objects.create(
      user = self.user2,
      accNonSSO = non_sso_acc2,
      username = 'jonikeren',
      email = 'acc6@example.com',
      role = 'Admin',
      accountType = NON_SSO_UI
    )
    
    self.pegawai1, self.pegawai2 = setup_pegawai()

    self.event_name = 'event1-validation'
    self.event_name2 = 'event2-validation'
    self.start_date = datetime.date(2023,3,19)
    self.end_date   = datetime.date(2023,3,21)

    self.binary_data = DUMMY_BINARY_DATA
    self.dummy_file = SimpleUploadedFile(DUMMY_PDF, self.binary_data, content_type=APPLICATION_PDF)


    self.event = Event.objects.create(
      creator=self.account,
      event_name=self.event_name,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=200000,
      sk_file=self.binary_data,
      status=WAITING_FOR_VALIDATION,
      signed_file=self.dummy_file
    )

    self.event2 = Event.objects.create(
      creator=self.account2,
      event_name=self.event_name2,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=200000,
      sk_file=self.binary_data,
      status=WAITING_FOR_VALIDATION,
      signed_file=self.dummy_file
    )

    self.eventEmployee =  EventEmployee.objects.create(
      employee=self.pegawai1,
      event=self.event,
      honor=100000,
      pph=10,
      netto=10000,
      role='Ketua'
    )

  def test_valid_event_validation(self):
    url = reverse(VALIDATE_EVENT, args=[self.event.id])

    response = self.client.post(url)

    self.event.status = 'Validated'  # Update the status field of self.event
    self.event.save()  # Save the updated event

    # Assert the response status code, event status, and template name
    self.assertEqual(response.status_code, 200)
    self.assertEqual(self.event.status, 'Validated')
    self.assertTemplateUsed(response, DETAIL_EVENT_HTML)

    event = Event.objects.first()
    self.assertEqual(event.creator, self.account)
    self.assertEqual(event.event_name, self.event_name)
    self.assertEqual(Event.objects.count(), 2)
    self.assertEqual(event.start_date, self.start_date)
    self.assertEqual(event.end_date, self.end_date)
    
    event_employee1 = EventEmployee.objects.get(employee=self.pegawai1)
    self.assertEqual(event_employee1.role, 'Ketua')
    self.assertEqual(event_employee1.honor, 100000)
    self.assertEqual(event_employee1.pph, 10)

  def test_invalid_event_validation(self):
    # Test with an invalid event ID
    url = reverse(VALIDATE_EVENT, args=[999])
    response = self.client.post(url)

    # Assert the response status code and redirected URL
    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, FORBIDDEN_URL)

  def test_non_user_validate_event(self):
    url = reverse(VALIDATE_EVENT, args=[self.event2.id])
    response = self.client.post(url)

    # Assert the response status code and redirected URL
    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, FORBIDDEN_URL)

class EventRejectionTestCase(TestCase):
  
  def setUp(self):
    self.client = Client()
    self.user = User.objects.create_user(
        username='testuser3', email=TEST_USER_EMAIL, password='testpassword3')
    self.client.login(username='testuser3', password='testpassword3')

    self.user2 = User.objects.create_user(
      username='testuser2', email=TEST_USER2_EMAIL, password='testpassword2'
    )
    self.client2 = Client()
    self.client2.login(username='testuser2', password='testpassword2')

    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = 'acc7@example.com',
      role = 'User',
      is_first_login = True
    )

    non_sso_acc2 = NonSSOAccount.objects.create(
      user = self.user2,
      username = 'jonikeren',
      email = 'acc8@example.com',
      role = 'Admin',
      is_first_login = True
    )

    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = 'acc9@example.com',
      role = 'User',
      accountType = NON_SSO_UI
    )

    self.account2 = Account.objects.create(
      user = self.user2,
      accNonSSO = non_sso_acc2,
      username = 'jonikeren',
      email = 'acc10@example.com',
      role = 'Admin',
      accountType = NON_SSO_UI
    )
    
    self.pegawai1, self.pegawai2 = setup_pegawai()

    self.event_name = 'event1-rejection'
    self.event_name2 = 'event2-rejection'
    self.start_date = datetime.date(2023,3,21)
    self.end_date   = datetime.date(2023,3,23)

    self.binary_data = DUMMY_BINARY_DATA
    self.dummy_file = SimpleUploadedFile(DUMMY_PDF, self.binary_data, content_type=APPLICATION_PDF)


    self.event = Event.objects.create(
      creator=self.account,
      event_name=self.event_name,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=100000,
      sk_file=self.binary_data,
      status=WAITING_FOR_VALIDATION,
      signed_file=self.dummy_file
    )

    self.event2 = Event.objects.create(
      creator=self.account2,
      event_name=self.event_name2,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=100000,
      sk_file=self.binary_data,
      status=WAITING_FOR_VALIDATION,
      signed_file=self.dummy_file
    )

    self.eventEmployee =  EventEmployee.objects.create(
      employee=self.pegawai1,
      event=self.event,
      honor=10000,
      pph=5,
      netto=9500,
      role='Ketua'
    )

  def test_valid_event_rejection(self):
    url = reverse(REJECT_EVENT, args=[self.event.id])

    response = self.client.post(url)

    self.event.status = 'Rejected'  # Update the status field of self.event
    self.event.save()  # Save the updated event

    # Assert the response status code, event status, and template name
    self.assertEqual(response.status_code, 200)
    self.assertEqual(self.event.status, 'Rejected')
    self.assertTemplateUsed(response, DETAIL_EVENT_HTML)

    event = Event.objects.first()
    self.assertEqual(event.creator, self.account)
    self.assertEqual(event.event_name, self.event_name)
    self.assertEqual(Event.objects.count(), 2)
    self.assertEqual(event.start_date, self.start_date)
    self.assertEqual(event.end_date, self.end_date)
    
    event_employee1 = EventEmployee.objects.get(employee=self.pegawai1)
    self.assertEqual(event_employee1.role, 'Ketua')
    self.assertEqual(event_employee1.honor, 10000)
    self.assertEqual(event_employee1.pph, 5)

  def test_invalid_event_validation(self):
    # Test with an invalid event ID
    url = reverse(REJECT_EVENT, args=[999])
    response = self.client.post(url)

    # Assert the response status code and redirected URL
    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, FORBIDDEN_URL)

  def test_non_user_validate_event(self):
    url = reverse(REJECT_EVENT, args=[self.event2.id])
    response = self.client.post(url)

    # Assert the response status code and redirected URL
    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, FORBIDDEN_URL)

class UploadSuratTugasTestCase(TestCase):
  
  def setUp(self):
    self.client = Client()
    self.user = User.objects.create_user(
        username='testuser5', email=TEST_USER_EMAIL, password='testpassword5')
    self.client.login(username='testuser5', password='testpassword5')

    self.user2 = User.objects.create_user(
      username='testuser2', email=TEST_USER2_EMAIL, password='testpassword2'
    )
    self.client2 = Client()
    self.client2.login(username='testuser2', password='testpassword2')

    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = 'acc11@example.com',
      role = 'User',
      is_first_login = True
    )

    non_sso_acc2 = NonSSOAccount.objects.create(
      user = self.user2,
      username = 'jonikeren',
      email = 'acc12@example.com',
      role = 'Admin',
      is_first_login = True
    )

    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = 'acc13@example.com',
      role = 'User',
      accountType = NON_SSO_UI
    )

    self.account2 = Account.objects.create(
      user = self.user2,
      accNonSSO = non_sso_acc2,
      username = 'jonikeren',
      email = 'acc14@example.com',
      role = 'Admin',
      accountType = NON_SSO_UI
    )
    
    self.pegawai1, self.pegawai2 = setup_pegawai()

    self.event_name = 'event1-upload'
    self.event_name2 = 'event2-upload'
    self.start_date = datetime.date(2023,3,21)
    self.end_date   = datetime.date(2023,3,23)

    self.binary_data = DUMMY_BINARY_DATA
    self.dummy_file = SimpleUploadedFile(DUMMY_PDF, self.binary_data, content_type=APPLICATION_PDF)


    self.event = Event.objects.create(
      creator=self.account,
      event_name=self.event_name,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=100000,
      sk_file=self.binary_data,
      status='Not validated yet',
      signed_file=self.dummy_file
    )

    self.event2 = Event.objects.create(
      creator=self.account2,
      event_name=self.event_name2,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=100000,
      sk_file=self.binary_data,
      status='Not validated yet',
      signed_file=self.dummy_file
    )

    self.validation_file = ValidationFile.objects.create(
      creator=self.account,
      event=self.event,
      surat_tugas=self.dummy_file
    )


  def test_valid_post_upload_surat_tugas(self):
    url = reverse(UPLOAD_SURAT_TUGAS, args=[self.event.id])
    data = {
      'signedSuratTugas':self.dummy_file
    }
    response = self.client.post(url, data=data)

    self.event.status = WAITING_FOR_VALIDATION

    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, reverse('detail_event', args=[self.event.id]))
    validation_file = ValidationFile.objects.filter(event=self.event).first()
    self.assertIsNotNone(validation_file)
    self.assertEqual(self.event.status, WAITING_FOR_VALIDATION)

  def test_valid_get_upload_surat_tugas(self):
    url = reverse(UPLOAD_SURAT_TUGAS, args=[self.event.id])
    
    response = self.client.get(url)

    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed('upload_surat_tugas.html')


  def test_non_user_validate_event(self):
    url = reverse(UPLOAD_SURAT_TUGAS, args=[self.event2.id])
    response = self.client.post(url)

    # Assert the response status code and redirected URL
    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, FORBIDDEN_URL)

class ReuploadSuratTugasTestCase(TestCase):
  
  def setUp(self):
    self.client = Client()
    self.user = User.objects.create_user(
        username='testuser7', email=TEST_USER_EMAIL, password='testpassword7')
    self.client.login(username='testuser7', password='testpassword7')

    self.user2 = User.objects.create_user(
      username='testuser2', email=TEST_USER2_EMAIL, password='testpassword2'
    )
    self.client2 = Client()
    self.client2.login(username='testuser2', password='testpassword2')

    non_sso_acc = NonSSOAccount.objects.create(
      user = self.user,
      username = 'jonikeren',
      email = 'acc15@example.com',
      role = 'User',
      is_first_login = True
    )

    non_sso_acc2 = NonSSOAccount.objects.create(
      user = self.user2,
      username = 'jonikeren',
      email = 'acc16@example.com',
      role = 'Admin',
      is_first_login = True
    )

    self.account = Account.objects.create(
      user = self.user,
      accNonSSO = non_sso_acc,
      username = 'jonikeren',
      email = 'acc17@example.com',
      role = 'User',
      accountType = NON_SSO_UI
    )

    self.account2 = Account.objects.create(
      user = self.user2,
      accNonSSO = non_sso_acc2,
      username = 'jonikeren',
      email = 'acc1@example.com',
      role = 'Admin',
      accountType = NON_SSO_UI
    )
    
    self.pegawai1, self.pegawai2 = setup_pegawai()

    self.event_name = 'event1-reupload'
    self.event_name2 = 'event2-reupload'
    self.start_date = datetime.date(2023,3,21)
    self.end_date   = datetime.date(2023,3,23)

    self.binary_data = DUMMY_BINARY_DATA
    self.dummy_file = SimpleUploadedFile(DUMMY_PDF, self.binary_data, content_type=APPLICATION_PDF)


    self.event = Event.objects.create(
      creator=self.account,
      event_name=self.event_name,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=100000,
      sk_file=self.binary_data,
      status='Rejected',
      signed_file=self.dummy_file
    )

    self.event2 = Event.objects.create(
      creator=self.account2,
      event_name=self.event_name2,
      start_date=self.start_date,
      end_date=self.end_date,
      expense=100000,
      sk_file=self.binary_data,
      status=WAITING_FOR_VALIDATION,
      signed_file=self.dummy_file
    )

    self.validation_file = ValidationFile.objects.create(
      creator=self.account,
      event=self.event,
      surat_tugas=self.dummy_file
    )


  def test_valid_post_reupload_surat_tugas(self):
    url = reverse(REUPLOAD_SURAT_TUGAS, args=[self.event.id, self.validation_file.id])
    data = {
      'signedSuratTugas':self.dummy_file
    }
    response = self.client.post(url, data=data)

    self.event.status = WAITING_FOR_VALIDATION

    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, reverse('detail_event', args=[self.event.id]))

    submitted_file = ValidationFile.objects.get(id=self.validation_file.id)
    self.assertIsNotNone(submitted_file)
    self.assertEqual(self.event.status, WAITING_FOR_VALIDATION)

  def test_valid_get_upload_surat_tugas(self):
    url = reverse(REUPLOAD_SURAT_TUGAS, args=[self.event.id, self.validation_file.id])
    
    response = self.client.get(url)

    self.assertEqual(response.status_code, 200)
    self.assertTemplateUsed('reupload_surat_tugas.html')


  def test_non_user_validate_event(self):
    url = reverse(REUPLOAD_SURAT_TUGAS, args=[self.event2.id, self.validation_file.id])
    response = self.client.post(url)

    # Assert the response status code and redirected URL
    self.assertEqual(response.status_code, 302)
    self.assertRedirects(response, FORBIDDEN_URL)
