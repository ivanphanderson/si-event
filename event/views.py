from django.shortcuts import render, redirect
from .models import Event, EventEmployee, ValidationFile
from pegawai.models import Pegawai
from pegawai.views import add_log
from account.models import Account
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q
from django.http import JsonResponse
from .validators import validate_event_employee_fields
from django.views.decorators.http import require_http_methods
from django.db.models import Sum
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.shared import Pt
import docx
from io import BytesIO
from django.http import HttpResponse
from django.urls import reverse



CREATE_EVENT = "create_event.html"
EVENT_LIST = "event_list.html"
FORBIDDEN_URL = "home:forbidden"
LOGIN_URL = "authentication:login"
LIST_NUMBER = 'List Number'
DETAIL_EVENT_HTML = "detail_event.html"

def get_event_data(event):
    event_employees = EventEmployee.objects.filter(event=event)
    event_emps = EventEmployee.objects.all().filter(event=event.id)
    pph_in_rp = 0
    for emp in event_emps:
        pph_in_rp += emp.pph/100.0 * emp.honor
    total_bruto = event_emps.aggregate(Sum("honor"))["honor__sum"]
    total_pph = int(pph_in_rp)
    total_netto = event_emps.aggregate(Sum("netto"))["netto__sum"]

    return event_employees, total_bruto, total_pph, total_netto

@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET", "POST"])
def create_event(request):
    if request.method == "POST":
        body = request.POST

        event_name = body.get("event_name")
        start_date = body.get("start_date")
        end_date = body.get("end_date")
        action = body.get("action")

        if action == "add_roles":
            request.session["event_name"] = event_name
            request.session["start_date"] = start_date
            request.session["end_date"] = end_date
            return render(request, "input_employee.html")
        else:
            account = Account.objects.get(user=request.user)
            Event.objects.create(
                creator=account,
                event_name=event_name,
                start_date=start_date,
                end_date=end_date,
            )
            action = "Create " + event_name + " event"
            add_log(account, action)

            request.session.pop("event_name", None)
            request.session.pop("start_date", None)
            request.session.pop("end_date", None)
            return get_events(request, action)
    account = Account.objects.get(user=request.user)
    if account.role == "User":
        return render(request, CREATE_EVENT)
    else:
        return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["POST"])
def input_employee_to_event(request):
    account = Account.objects.get(user=request.user)
    form_data = request.POST

    if form_data["num_fields"] != "":
        num_fields = int(form_data["num_fields"])
    else:
        num_fields = 0

    total_honor = 0
    for idx in range(num_fields):
        if form_data[f"honor_field_{idx}"] != "":
            total_honor += abs(int(form_data[f"honor_field_{idx}"]))

    new_event = Event.objects.create(
        creator=account,
        event_name=request.session["event_name"],
        start_date=request.session["start_date"],
        end_date=request.session["end_date"],
        expense=total_honor,
    )

    action = "Create " + request.session["event_name"] + " event"
    add_log(Account.objects.get(user=request.user), action)

    for idx in range(num_fields):
        if f"dropdown-select_{idx}" in form_data:
            employee_no = form_data[f"dropdown-select_{idx}"]
            role = form_data[f"role_field_{idx}"]
            honor = form_data[f"honor_field_{idx}"]
            pph = form_data[f"pph_field_{idx}"]

            role, pph, honor = validate_event_employee_fields(role, pph, honor, idx)
            if str(employee_no).isnumeric():
                pegawai = Pegawai.objects.get(employee_no=employee_no)
                EventEmployee.objects.create(
                    employee=pegawai, event=new_event, honor=honor, pph=pph, role=role
                )

    return get_events(request, action)


@require_http_methods(["GET"])
def get_options(request):
    search_term = request.GET.get("search", "")
    employees = Pegawai.objects.filter(
        Q(employee_no__icontains=search_term) | Q(employee_name__icontains=search_term)
    )
    employees = [
        {"employee_no": e.employee_no, "employee_name": e.employee_name}
        for e in employees
    ]
    return JsonResponse(employees, safe=False)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET", "POST"])
def get_events(request, success_message=None):
    account = Account.objects.get(user=request.user)
    event_data = Event.objects.all().order_by("event_name")
    owner_data = dict()

    for event in event_data:
        if event.creator == account:
            owner_data[f"{event.event_name}"] = True

    if success_message is not None:
        return render(
            request,
            EVENT_LIST,
            {
                "event_data": event_data,
                "owner_data": owner_data,
                "success_message": success_message,
                "role": account.role,
            },
        )

    return render(
        request,
        EVENT_LIST,
        {"event_data": event_data, "owner_data": owner_data, "role": account.role},
    )


@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET"])
def riwayat_events(request):
    account = Account.objects.get(user=request.user)
    if account.role == "User":
        event_data = Event.objects.filter(creator=account).order_by("event_name")
        return render(
            request,
            "riwayat_event.html",
            {"event_data": event_data, "role": account.role},
        )
    else:
        return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET","POST"])
def detail_event(request, id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["account"] = account
    context["role"] = account.role

    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account and account.role == "User":
            return redirect(FORBIDDEN_URL)

        event_employees = EventEmployee.objects.filter(event=event)
        context["event"] = event
        context["event_employees"] = event_employees

        event_emps = EventEmployee.objects.all().filter(event=event.id)
        pph_in_rp = 0
        for emp in event_emps:
            pph_in_rp += emp.pph/100.0 * emp.honor
        context["total_bruto"] = event_emps.aggregate(Sum("honor"))["honor__sum"]
        context["total_pph"] = int(pph_in_rp)
        context["total_netto"] = event_emps.aggregate(Sum("netto"))["netto__sum"]

        if event.signed_file:
            submitted_file = ValidationFile.objects.filter(event=event).first()
            context['file_id'] = submitted_file.id

        return render(request, DETAIL_EVENT_HTML, context)
    return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET"])
def update_event(request, id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["role"] = account.role
    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account or event.status == 'Validated':
            return redirect(FORBIDDEN_URL)
        context["event"] = event
        return render(request, "update_event.html", context)
    return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["POST", "GET"])
def delete_event(request, id):
    account = Account.objects.get(user=request.user)
    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account:
            return redirect(FORBIDDEN_URL)
        Event.objects.filter(id=id).delete()
        return redirect('get_events')
    return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["POST"])
def submit_update_event(request, id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["role"] = account.role
    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account or event.status == 'Validated':
            return redirect(FORBIDDEN_URL)
        body = request.POST
        event.event_name = body.get("event_name")
        event.start_date = body.get("start_date")
        event.end_date = body.get("end_date")
        event.save()

        action = f"Updated {event.event_name} event"
        add_log(Account.objects.get(user=request.user), action)
        messages.success(request, "Event is updated successfully.")
        return redirect(f"/event/detail/{id}")
    return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET"])
def input_employee_to_existing_event(request, id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["role"] = account.role
    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account or event.status == 'Validated':
            return redirect(FORBIDDEN_URL)
        context["event"] = event
        return render(request, "input_employee_to_existing_event.html", context)
    return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["POST"])
def submit_input_employee_to_existing_event(request, id):
    account = Account.objects.get(user=request.user)
    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account or event.status == 'Validated':
            return redirect(FORBIDDEN_URL)
        form_data = request.POST

        if form_data["num_fields"] != "":
            num_fields = int(form_data["num_fields"])
        else:
            num_fields = 0

        total_honor = calculate_total_honor(form_data, num_fields)

        event.expense += total_honor
        event.save()

        action = f"Added Employee to {event.event_name} event"
        add_log(Account.objects.get(user=request.user), action)

        for idx in range(num_fields):
            if f"dropdown-select_{idx}" in form_data:
                process_employee_data(event, form_data, idx)

        messages.success(request, "Employees added successfully.")
        return redirect(f"/event/detail/{id}")
    return redirect(FORBIDDEN_URL)


def calculate_total_honor(form_data, num_fields):
    total_honor = 0
    for idx in range(num_fields):
        honor_field = form_data.get(f"honor_field_{idx}")
        if honor_field:
            total_honor += abs(int(honor_field))
    return total_honor


def process_employee_data(event, form_data, idx):
    employee_no = form_data[f"dropdown-select_{idx}"]
    role = form_data[f"role_field_{idx}"]
    honor = form_data[f"honor_field_{idx}"]
    pph = form_data[f"pph_field_{idx}"]
    role, pph, honor = validate_event_employee_fields(role, pph, honor, idx)
    if str(employee_no).isnumeric():
        pegawai = Pegawai.objects.get(employee_no=employee_no)
        EventEmployee.objects.create(
            employee=pegawai, event=event, honor=honor, pph=pph, role=role
        )


@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET"])
def update_event_employee_by_id(request, id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["role"] = account.role
    if id.isdigit() and EventEmployee.objects.filter(id=id).first():
        event_employee = EventEmployee.objects.get(id=id)
        context["event_employee"] = event_employee
        if event_employee.event.creator != account or event_employee.event.status == 'Validated':
            return redirect(FORBIDDEN_URL)
        return render(request, "update_event_employee.html", context)
    return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["POST"])
def update_employee_to_event_by_id(request, id):
    account = Account.objects.get(user=request.user)
    form_data = request.POST
    if id.isdigit() and EventEmployee.objects.filter(id=id).first():
        event_employee = EventEmployee.objects.get(id=id)
        event = event_employee.event

        if event.creator != account or event.status == 'Validated':
            return redirect(FORBIDDEN_URL)

        employee_no = form_data["dropdown-select_0"]
        role = form_data["role_field_0"]
        honor = form_data["honor_field_0"]
        pph = form_data["pph_field_0"]

        role, pph, honor = validate_event_employee_fields(role, pph, honor, 0)
        if str(employee_no).isdigit():
            pegawai = Pegawai.objects.get(employee_no=employee_no)

            event.expense = event.expense - event_employee.honor + honor
            event.save()

            event_employee.employee = pegawai
            event_employee.role = role
            event_employee.honor = honor
            event_employee.pph = pph
            event_employee.save()

        action = f"Updated {event_employee.employee.employee_name} in {event.event_name} event"
        add_log(Account.objects.get(user=request.user), action)

        messages.success(
            request, f"{event_employee.employee.employee_name} is updated successfully."
        )
        return redirect(f"/event/detail/{event.id}")
    return redirect(FORBIDDEN_URL)


@login_required(login_url=LOGIN_URL)
@require_http_methods(["POST"])
def delete_event_employee_by_id(request, id):
    account = Account.objects.get(user=request.user)
    if id.isdigit() and EventEmployee.objects.filter(id=id).first():
        event_employee = EventEmployee.objects.get(id=id)
        event = event_employee.event
        if event.creator != account or event.status == 'Validated':
            return redirect(FORBIDDEN_URL)
        event_employee.delete()

        action = f"Deleted {event_employee.employee.employee_name} from {event.event_name} event"
        add_log(Account.objects.get(user=request.user), action)

        messages.success(
            request, f"{event_employee.employee.employee_name} is deleted successfully."
        )
        return redirect(f"/event/detail/{event.id}")
    return redirect(FORBIDDEN_URL)

class GenerateDocs:
    def set_font(self, run, bold=False, underline=False):
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = bold
        font.underline = WD_UNDERLINE.SINGLE if underline else WD_UNDERLINE.NONE

    def set_paragraph_format(self, paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, left_indent=0):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment
        paragraph_format.left_indent = Pt(left_indent)

    def create_docs(self, body, event_id):
        document = Document()
        start_date = self.get_start_date(event_id)
        employee = self.get_event_employee(event_id)
        nomor_surat_tugas = body.get("nomor_surat_tugas")
        nama_pj = body.get("nama_pj")
        jabatan_pj = body.get("jabatan_pj")
        perihal_event = body.get("perihal_event")
        tugas_panitia = body.get("tugas_panitia")
        target_anggaran = body.get("target_anggaran")

        paragraph1 = document.add_paragraph("SURAT TUGAS")
        self.set_font(paragraph1.runs[0], True, True)
        self.set_paragraph_format(paragraph1, WD_ALIGN_PARAGRAPH.CENTER)

        paragraph2 = document.add_paragraph(f"No.: ST-{nomor_surat_tugas}/UN2.F11.D/HKP.02.04/2019")
        self.set_font(paragraph2.runs[0])
        self.set_paragraph_format(paragraph2, WD_ALIGN_PARAGRAPH.CENTER)

        document.add_paragraph()

        paragraph3 = document.add_paragraph("Yang bertanda tangan di bawah ini:")
        self.set_font(paragraph3.runs[0])

        paragraph4 = document.add_paragraph(f"Nama: {nama_pj}")
        self.set_font(paragraph4.runs[0])
        self.set_paragraph_format(paragraph4, left_indent=36)

        run = paragraph4.add_run() 
        run.add_break(docx.text.run.WD_BREAK.LINE)
        run2 = paragraph4.add_run(f'Jabatan: {jabatan_pj}')
        self.set_font(run2)

        paragraph6 = document.add_paragraph(f"dengan ini menugaskan kepada nama-nama staf dan karyawan terlampir untuk menjadi Panitia {perihal_event}")
        self.set_font(paragraph6.runs[0])
        self.set_paragraph_format(paragraph6)

        paragraph7 = document.add_paragraph(f"Jadwal Kegiatan: {start_date}", style = LIST_NUMBER)
        self.set_font(paragraph7.runs[0])

        paragraph8 = document.add_paragraph("Tugas Panitia", style = LIST_NUMBER)
        self.set_font(paragraph8.runs[0])

        list_tugas_panitia = tugas_panitia.split('\n')

        for i in range(len(list_tugas_panitia)):
            paragraph8_3 = document.add_paragraph(f"{list_tugas_panitia[i].strip()}", style = 'List Number 2')
            self.set_font(paragraph8_3.runs[0])
        
        paragraph9 = document.add_paragraph(f"Pengeluaran biaya yang ditimbulkan akibat pemberlakuan Surat Tugas ini dibebankan secara proporsional pada {target_anggaran}", style = LIST_NUMBER)
        self.set_font(paragraph9.runs[0])

        paragraph10 = document.add_paragraph(f"Surat Tugas ini berlaku sejak tanggal ditetapkan sampai berakhirnya kegiatan {perihal_event}", style = LIST_NUMBER)
        self.set_font(paragraph10.runs[0])

        paragraph11 = document.add_paragraph("Demikian Surat Tugas ini dibuat untuk dilaksanakan dengan penuh tanggung jawab. Apabila di kemudian hari ternyata terdapat kekeliruan dalam Surat Tugas ini, akan diadakan perbaikan seperlunya.")
        self.set_font(paragraph11.runs[0])

        document.add_paragraph()

        paragraph12 = document.add_paragraph("Ditetapkan di	:	Jakarta")
        self.set_font(paragraph12.runs[0])
        self.set_paragraph_format(paragraph12, left_indent=216)

        run12 = paragraph12.add_run() 
        run12.add_break(docx.text.run.WD_BREAK.LINE)
        run12_1 = paragraph12.add_run('Pada Tanggal	: Tanggal   Bulan   Tahun')
        self.set_font(run12_1)

        paragraph13 = document.add_paragraph(f"{jabatan_pj}")
        self.set_font(paragraph13.runs[0])
        self.set_paragraph_format(paragraph13, left_indent=216)

        document.add_paragraph()
        document.add_paragraph()

        paragraph14 = document.add_paragraph(f"{nama_pj}")
        self.set_font(paragraph14.runs[0], True)
        self.set_paragraph_format(paragraph14, left_indent=216)

        document.add_page_break()

        paragraph1_2 = document.add_paragraph("Lampiran Surat Tugas Dekan Fakultas Ilmu Komputer Universitas Indonesia")
        self.set_font(paragraph1_2.runs[0])

        paragraph2_2 = document.add_paragraph(f"No.: ST-{nomor_surat_tugas}/UN2.F11.D/HKP.02.04/2019")
        self.set_font(paragraph2_2.runs[0])

        paragraph3_2 = document.add_paragraph(f"Perihal	:   {perihal_event}")
        self.set_font(paragraph3_2.runs[0])

        employee_dict = {}
        for emp in employee:
            if emp[0] not in employee_dict:
                employee_dict[emp[0]] = []
            employee_dict[emp[0]].append(emp[1])
        
        for role, employees in employee_dict.items():
            paragraph = document.add_paragraph(f"{role}      :       {employees[0]}", style="List Bullet")
            self.set_font(paragraph.runs[0])
            self.set_paragraph_format(paragraph, left_indent=36)
            for i in range(1, len(employees)):
                run = paragraph.add_run()
                run.add_break(docx.text.run.WD_BREAK.LINE)
                run.add_text(f"{employees[i]}")
                self.set_font(run)
                
        document.add_paragraph()

        paragraph9_2 = document.add_paragraph("Ditetapkan di	:	Jakarta")
        self.set_font(paragraph9_2.runs[0])
        self.set_paragraph_format(paragraph9_2, left_indent=216)

        run12 = paragraph9_2.add_run() 
        run12.add_break(docx.text.run.WD_BREAK.LINE)
        run12_1 = paragraph9_2.add_run('Pada Tanggal	: Tanggal   Bulan   Tahun')
        self.set_font(run12_1)

        paragraph10_2 = document.add_paragraph(f"{jabatan_pj}")
        self.set_font(paragraph10_2.runs[0])
        self.set_paragraph_format(paragraph10_2, left_indent=216)

        document.add_paragraph()
        document.add_paragraph()

        paragraph11_2 = document.add_paragraph(f"{nama_pj}")
        self.set_font(paragraph11_2.runs[0], True)
        self.set_paragraph_format(paragraph11_2, left_indent=216)

        return document

    def download_file(self, body, event_id):
        # Create a new Word document
        document = self.create_docs(body, event_id)

        # Create a BytesIO object to write the document to
        output = BytesIO()
        document.save(output)
        output.seek(0)

        # Create an HttpResponse object with the correct MIME type
        response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=SURAT TUGAS.docx'

        return response
    
    def get_event_name(self, event_id):
        if Event.objects.filter(id=event_id).first():
            event = Event.objects.get(id=event_id)
            return event.event_name
        
    def get_event_employee(self, event_id):
        event = Event.objects.get(id=event_id)
        event_employees = EventEmployee.objects.filter(event=event)
        lst = []
        for event_employee in event_employees:
            lst.append([ event_employee.role, event_employee.employee.employee_name])
        return lst
    
    def get_start_date(self, event_id):
        if Event.objects.filter(id=event_id).first():
            event = Event.objects.get(id=event_id)
            return event.start_date
        
    def get_end_date(self, event_id):
        if Event.objects.filter(id=event_id).first():
            event = Event.objects.get(id=event_id)
            return event.end_date

def generate_docs(body, event_id):
    generate_docs_file = GenerateDocs()

    # Call the download_file method to generate and download the document
    response = generate_docs_file.download_file(body, event_id)

    # Return the response object to the user
    return response

@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET", "POST"])
def form_surat_tugas(request, event_id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["account"] = account
    context["role"] = account.role

    if request.method == 'POST':
        body = request.POST
        return generate_docs(body, event_id)

    if Event.objects.filter(id=event_id).first():
        event = Event.objects.get(id=event_id)
        context["event"] = event

        if event.creator != account:
            return redirect(FORBIDDEN_URL)

        return render(request, 'form_surat_tugas.html', context)
    
    return redirect(FORBIDDEN_URL)

@require_http_methods(["GET","POST"])
def validate_event(request, id):
    
    context = {}
    account = Account.objects.get(user=request.user)
    context["account"] = account
    context["role"] = account.role

    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account and account.role == "User":
            return redirect(FORBIDDEN_URL)

        context["event"] = event
        event.status = 'Validated'
        event.save()

        event_employees, total_bruto, total_pph, total_netto = get_event_data(event)

        context["event_employees"] = event_employees
        context["total_bruto"] = total_bruto
        context["total_pph"] = total_pph
        context["total_netto"] = total_netto

        return render(request, DETAIL_EVENT_HTML, context)
    return redirect(FORBIDDEN_URL)

@login_required(login_url=LOGIN_URL)
@require_http_methods(["POST"])
def reject_event(request, id):

    context = {}
    account = Account.objects.get(user=request.user)
    context["account"] = account
    context["role"] = account.role

    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account and account.role == "User":
            return redirect(FORBIDDEN_URL)

        context["event"] = event
        event.status = 'Rejected'
        rejection_reason = request.POST.get('rejection_reason')
        event.rejection_reason = rejection_reason
        event.save()

        event_employees, total_bruto, total_pph, total_netto = get_event_data(event)
        
        context["event_employees"] = event_employees
        context["total_bruto"] = total_bruto
        context["total_pph"] = total_pph
        context["total_netto"] = total_netto

        return render(request, DETAIL_EVENT_HTML, context)
    return redirect(FORBIDDEN_URL)

@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET", "POST"])
def upload_surat_tugas(request, id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["account"] = account
    context["role"] = account.role

    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account and account.role == "User":
            return redirect(FORBIDDEN_URL)

        context["event"] = event

        if request.method == 'POST':
            signed_file = request.FILES.get('signedSuratTugas')
            validation_file = ValidationFile(
                creator = account,
                event=event,
                surat_tugas = signed_file
            )
            validation_file.save()
            event.signed_file = signed_file
            event.status = 'Waiting for validation'
            event.save()

            return redirect(reverse('detail_event', args=[event.id]))

        return render(request, 'upload_surat_tugas.html', context)

@login_required(login_url=LOGIN_URL)
@require_http_methods(["GET", "POST"])
def reupload_surat_tugas(request, id, file_id):
    context = {}
    account = Account.objects.get(user=request.user)
    context["account"] = account
    context["role"] = account.role

    if id.isdigit() and Event.objects.filter(id=id).first():
        event = Event.objects.get(id=id)
        if event.creator != account and account.role == "User":
            return redirect(FORBIDDEN_URL)

        context["event"] = event
        submitted_file = ValidationFile.objects.filter(id=file_id).first()
        context['file_id'] = submitted_file.id

        if request.method == 'POST':
            signed_file = request.FILES.get('signedSuratTugas')

            submitted_file.surat_tugas = signed_file

            submitted_file.save()
            event.signed_file = signed_file
            event.status = 'Waiting for validation'
            event.rejection_reason = ''
            event.save()

            return redirect(reverse('detail_event', args=[event.id]))

        return render(request, 'reupload_surat_tugas.html', context)